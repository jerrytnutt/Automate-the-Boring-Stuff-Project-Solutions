# Write a program that will generate a Word document with custom invitations.
# Use a guests.txt file with a list of guest names, one name per line.
# There should be one invitation per page in the resulting Word document.

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def create_greeting_cards():
  # Append all guest names from word documents into a list
  doc = docx.Document('guest.docx')
  guest_list = []
  for paragraphs in doc.paragraphs:
    guest_list.append(paragraphs.text)
  length = len(guest_list)
  # Lines used for each custom invitation
  line_list = ['It would be the my pleasure to have the company of',
  'none','at 11010 memory lane on the evening of','April 1st',"at 7 O'clock"]

  doc = docx.Document()
  for i in range(length):
    line_list[1] = guest_list[i]
    for j in range(length):
      paragraph = doc.add_paragraph()
      line = paragraph.add_run(line_list[j])
      paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
      font = line.font
      # Assign the proper font and style to each line
      if j in [0,2,4]:
        font.bold = True
        font.name = 'brush script std'
        font.size = docx.shared.Pt(14)
      elif j == 1:
        font.bold = True
        font.size = docx.shared.Pt(24)
      else:
        font.size = docx.shared.Pt(20)
    doc.add_page_break()
  doc.save('greetingCards.docx')
 

if __name__ == "__main__":
  create_greeting_cards()
